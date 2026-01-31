#!/usr/bin/env python3
"""
Convert all .udf files under a directory (recursively) to .docx (Word) files.
Strategy for each file (in order):
 1) Try pandoc (best for text/markdown-like files)
 2) Try LibreOffice headless (best for proprietary formats that LibreOffice can open)
 3) Fallback: read as text and create a simple .docx with python-docx

Usage:
  python3 convert_udf_to_docx.py /path/to/UYAP-TOOL [--out-dir /path/to/out] [--overwrite] [--dry-run]

Dependencies (macOS):
  brew install pandoc
  brew install --cask libreoffice
  pip3 install python-docx

"""
import argparse
import subprocess
import tempfile
import shutil
from pathlib import Path
import sys
from docx import Document
import os


def find_udf_files(root: Path):
    return list(root.rglob('*.udf')) + list(root.rglob('*.UDF'))


def run_pandoc(in_path: Path, out_path: Path):
    try:
        subprocess.run(['pandoc', str(in_path), '-o', str(out_path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return True, 'pandoc'
    except FileNotFoundError:
        return False, 'pandoc-not-installed'
    except subprocess.CalledProcessError:
        return False, 'pandoc-failed'


def run_libreoffice(in_path: Path, out_path: Path):
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        return False, 'libreoffice-not-installed'
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run([soffice, '--headless', '--convert-to', 'docx', '--outdir', tmpdir, str(in_path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            return False, 'libreoffice-failed'
        # find converted file
        converted = Path(tmpdir) / (in_path.stem + '.docx')
        if converted.exists():
            out_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(converted), str(out_path))
            return True, 'libreoffice'
        return False, 'libreoffice-no-output'


import zipfile
import re


def extract_zip_content_to_docx(in_path: Path, out_path: Path):
    """If the .udf is a zip-like archive (has content.xml), extract readable
    text (CDATA inside <content> tag) and write to a .docx after sanitizing
    control characters.
    """
    try:
        if not zipfile.is_zipfile(in_path):
            return False, 'not-zip'
        with zipfile.ZipFile(in_path) as z:
            if 'content.xml' not in z.namelist():
                return False, 'no-content-xml'
            raw = z.read('content.xml').decode('utf-8', errors='replace')
            m = re.search(r'<content[^>]*>(.*?)</content>', raw, re.DOTALL)
            if not m:
                return False, 'no-content-tag'
            cdata = m.group(1)
            # strip CDATA markers if present
            cdata = re.sub(r'^<!\[CDATA\[', '', cdata)
            cdata = re.sub(r'\]\]>\s*$', '', cdata)
            # remove control characters that break XML/Docx
            cdata = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', cdata)
            # normalize Windows newlines
            cdata = cdata.replace('\r\n', '\n')
            doc = Document()
            for line in cdata.splitlines():
                doc.add_paragraph(line)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out_path))
            return True, 'zip-content'
    except Exception as e:
        return False, f'zip-failed:{e}'


def fallback_text_to_docx(in_path: Path, out_path: Path):
    try:
        # try read with utf-8, fallback to latin1
        try:
            text = in_path.read_text(encoding='utf-8')
        except Exception:
            text = in_path.read_text(encoding='latin-1')
    except Exception as e:
        return False, f'read-failed:{e}'
    try:
        # sanitize control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return True, 'fallback-docx'
    except Exception as e:
        return False, f'docx-save-failed:{e}'


def convert_file(in_path: Path, out_path: Path, overwrite=False, dry_run=False):
    if out_path.exists() and not overwrite:
        return False, 'exists'
    if dry_run:
        return True, 'dry-run'

    # try pandoc
    ok, reason = run_pandoc(in_path, out_path)
    if ok:
        return True, reason
    # try libreoffice
    ok, reason = run_libreoffice(in_path, out_path)
    if ok:
        return True, reason
    # try extracting zip/ODF style content.xml
    ok, reason = extract_zip_content_to_docx(in_path, out_path)
    if ok:
        return True, reason
    # fallback: create simple docx
    ok, reason = fallback_text_to_docx(in_path, out_path)
    if ok:
        return True, reason

    return False, reason


def main():
    parser = argparse.ArgumentParser(description='Convert .udf files to .docx recursively')
    parser.add_argument('input_dir', type=Path, help='Root folder to search for .udf files')
    parser.add_argument('--out-dir', type=Path, default=None, help='Optional output root. If omitted, puts .docx next to original files')
    parser.add_argument('--overwrite', action='store_true', help='Overwrite existing .docx files')
    parser.add_argument('--dry-run', action='store_true', help='Show what would be done without writing files')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    args = parser.parse_args()

    root = args.input_dir
    if not root.exists() or not root.is_dir():
        print('Hata: verilen klasör bulunamadı:', root)
        sys.exit(1)

    udf_files = find_udf_files(root)
    if not udf_files:
        print('Hiç .udf dosyası bulunamadı altında:', root)
        sys.exit(0)

    total = len(udf_files)
    success = 0
    skipped = 0
    failed = 0

    for p in udf_files:
        if args.out_dir:
            rel = p.relative_to(root)
            out_path = args.out_dir / rel.with_suffix('.docx')
        else:
            out_path = p.with_suffix('.docx')

        ok, reason = convert_file(p, out_path, overwrite=args.overwrite, dry_run=args.dry_run)
        if ok:
            if reason == 'exists':
                skipped += 1
                if args.verbose:
                    print(f'SKIP (exists): {p} -> {out_path}')
            elif reason == 'dry-run':
                print(f'DRY-RUN: {p} -> {out_path}')
            else:
                success += 1
                if args.verbose:
                    print(f'OK ({reason}): {p} -> {out_path}')
        else:
            failed += 1
            print(f'FAILED ({reason}): {p} -> {out_path}')

    print('---')
    print(f'Total: {total}, Converted: {success}, Skipped: {skipped}, Failed: {failed}')


if __name__ == '__main__':
    main()
